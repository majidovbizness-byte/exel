"""Barcha xizmatlar: Gemini AI, hujjatlar, materiallar, obuna."""
import asyncio, base64, datetime, hashlib, io, json, math, os, re, tempfile, uuid
from PIL import Image, ImageDraw, ImageFont
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

GEN = os.path.abspath(os.path.join(os.path.dirname(__file__), "generated"))
os.makedirs(GEN, exist_ok=True)

# ══════════════════════════════════════════════════════════════
# MATERIALLAR
# ══════════════════════════════════════════════════════════════

async def mat_by_code(s, org_id, code):
    from models import Material
    code = code.strip().lstrip("0") or "0"
    r = await s.execute(select(Material).where(Material.organization_id==org_id, Material.code==code))
    return r.scalar_one_or_none()

async def mat_search(s, org_id, q, limit=8):
    from models import Material
    q = q.strip().casefold()
    r = await s.execute(select(Material).where(Material.organization_id==org_id))
    return [m for m in r.scalars() if q in m.name.casefold() or q in m.code.casefold()][:limit]

async def mat_get(s, mid):
    from models import Material
    return await s.get(Material, mid)

# ══════════════════════════════════════════════════════════════
# RAQAMLASH
# ══════════════════════════════════════════════════════════════

async def next_number(s, org_id):
    from models import NakCounter
    year = datetime.datetime.now().year
    r = await s.execute(select(NakCounter).where(
        NakCounter.organization_id==org_id, NakCounter.year==year).with_for_update())
    c = r.scalar_one_or_none()
    if c is None:
        c = NakCounter(organization_id=org_id, year=year, last_num=0)
        s.add(c); await s.flush()
    c.last_num += 1; await s.flush()
    return f"{c.last_num:04d}/{year}"

# ══════════════════════════════════════════════════════════════
# AUDIT
# ══════════════════════════════════════════════════════════════

async def audit_log(s, user_id, org_id, action, detail="", suspicious=False):
    from models import AuditLog
    s.add(AuditLog(user_id=user_id, org_id=org_id, action=action, detail=detail, suspicious=suspicious))
    await s.flush()

async def check_suspicious(s, user_id, org_id):
    from models import Nakladnaya
    warnings = []
    now = datetime.datetime.now(datetime.timezone.utc)
    if now.hour >= 22 or now.hour < 6:
        warnings.append(f"Tungi vaqtda ({now.strftime('%H:%M')}) nakladnoy yaratildi")
    hour_ago = now - datetime.timedelta(hours=1)
    r = await s.execute(select(Nakladnaya.id).where(
        Nakladnaya.creator_id==user_id, Nakladnaya.created_at>=hour_ago))
    if len(r.all()) >= 10:
        warnings.append("Bir soatda 10+ nakladnoy yaratildi")
    return warnings

# ══════════════════════════════════════════════════════════════
# OBUNA
# ══════════════════════════════════════════════════════════════

async def get_sub(s, org_id):
    from models import Subscription
    if not org_id: return None
    r = await s.execute(select(Subscription).where(Subscription.organization_id==org_id))
    return r.scalar_one_or_none()

async def get_plan_cfg(s, org_id, plan):
    from models import PlanConfig
    r = await s.execute(select(PlanConfig).where(PlanConfig.organization_id==org_id, PlanConfig.plan==plan))
    return r.scalar_one_or_none()

async def get_all_cfgs(s, org_id):
    from models import PlanConfig
    r = await s.execute(select(PlanConfig).where(PlanConfig.organization_id==org_id))
    return list(r.scalars())

async def current_cfg(s, org_id):
    sub = await get_sub(s, org_id)
    if not sub: return None
    return await get_plan_cfg(s, org_id, sub.plan)

async def ensure_default_plans(s, org_id):
    from models import PlanConfig, SubPlan
    existing = {c.plan for c in await get_all_cfgs(s, org_id)}
    defaults = [
        dict(plan=SubPlan.BASIC,    name="Asosiy",   price=150000, doc_limit=10,
             allow_excel=True,  allow_word=False, allow_pdf=False,
             allow_ocr=False,   allow_report=False, allow_transport=False,
             allow_calc=True,   allow_templates=2,  unlimited=False,
             allow_ai_calc=False, allow_ai_search=False, allow_ai_report=False,
             allow_ai_anomaly=False, allow_ai_briefing=False, allow_ai_voice=False,
             allow_ai_chat=False, allow_ai_translate=False),
        dict(plan=SubPlan.STANDARD, name="Standart", price=300000, doc_limit=100,
             allow_excel=True,  allow_word=True,  allow_pdf=False,
             allow_ocr=False,   allow_report=True,  allow_transport=True,
             allow_calc=True,   allow_templates=5,  unlimited=False,
             allow_ai_calc=True, allow_ai_search=True, allow_ai_report=True,
             allow_ai_anomaly=False, allow_ai_briefing=False, allow_ai_voice=False,
             allow_ai_chat=False, allow_ai_translate=False),
        dict(plan=SubPlan.PREMIUM,  name="Premium",  price=500000, doc_limit=999,
             allow_excel=True,  allow_word=True,  allow_pdf=True,
             allow_ocr=True,    allow_report=True,  allow_transport=True,
             allow_calc=True,   allow_templates=10, unlimited=True,
             allow_ai_calc=True, allow_ai_search=True, allow_ai_report=True,
             allow_ai_anomaly=True, allow_ai_briefing=True, allow_ai_voice=True,
             allow_ai_chat=True, allow_ai_translate=True),
    ]
    for d in defaults:
        if d["plan"] not in existing:
            s.add(PlanConfig(organization_id=org_id, **d))
    await s.flush()

# ══════════════════════════════════════════════════════════════
# EXCEL HUJJAT
# ══════════════════════════════════════════════════════════════

_T = Side(style="thin")
_BALL = Border(left=_T, right=_T, top=_T, bottom=_T)

def _c(ws, ref, val, sz=11, bold=False, align="left", wrap=False, border=False, bg=None):
    c = ws[ref]; c.value = val
    c.font = Font(name="Arial", size=sz, bold=bold)
    c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=wrap)
    if border: c.border = _BALL
    if bg: c.fill = PatternFill("solid", fgColor=bg)

def build_excel(nak, org_name, path, tpl=1):
    wb = Workbook(); ws = wb.active; ws.title = "Накладная"
    for col, w in {"A":6,"B":12,"C":44,"D":13,"E":14}.items():
        ws.column_dimensions[col].width = w
    _c(ws,"A1",f"№ {nak.number}",sz=10,bold=True)
    ws.merge_cells("A2:E2"); _c(ws,"A2","НАКЛАДНАЯ",sz=18,bold=True,align="center",
        bg="1F4E79" if tpl==3 else None)
    ws.merge_cells("A3:E3"); _c(ws,"A3","на перемещение товаров, материалов",sz=11,align="center")
    date_s = nak.created_at.strftime('"%d" %m. %Y') if nak.created_at else ""
    ws.merge_cells("D4:E4"); _c(ws,"D4",f"{date_s} год",sz=10,align="right")
    ws.merge_cells("A6:E6"); _c(ws,"A6",f"Организация: {org_name}",sz=11,bold=True)
    ws.row_dimensions[7].height=30
    ws.merge_cells("A7:E7"); _c(ws,"A7",f"Объект: {nak.object_text}",sz=11,wrap=True)
    ws.merge_cells("A9:E9");  _c(ws,"A9", f"Отправитель: {nak.sender or '________________________'}")
    ws.merge_cells("A10:E10");_c(ws,"A10",f"Получатель: {nak.receiver or '________________________'}")
    veh = f"{nak.vehicle_model} {nak.vehicle_plate}".strip() or "________________________"
    drv = nak.driver or "________________________"
    ws.merge_cells("A11:E11");_c(ws,"A11",f"Машина: {veh}   Водитель: {drv}",sz=10)
    dest = nak.destination or "________________________"
    ws.merge_cells("A12:E12");_c(ws,"A12",f"Куда: {dest}",sz=10,wrap=True)
    HR=13; ws.row_dimensions[HR].height=32
    hdr_bg = "2F75B6" if tpl in(3,7) else "D6E4F0" if tpl in(4,8) else "F2F2F2"
    hdr_fc = "FFFFFF" if tpl in(3,7) else "000000"
    for i,h in enumerate(["№\nп/п","Бух.счет","Наименование товар, материалов","Единица\nизмерения","Количество"]):
        col = get_column_letter(i+1)
        _c(ws,f"{col}{HR}",h,sz=10,bold=True,align="center",wrap=True,border=True,bg=hdr_bg)
        ws[f"{col}{HR}"].font = Font(name="Arial",size=10,bold=True,color=hdr_fc)
    DR=HR+1
    for r in range(24):
        er=DR+r; ws.row_dimensions[er].height=20
        alt = "F7FBFF" if r%2==0 and tpl==4 else None
        _c(ws,f"A{er}",r+1,align="center",border=True,bg=alt)
        it = nak.items[r] if r<len(nak.items) else None
        _c(ws,f"B{er}",it.code     if it else "",align="center",border=True,bg=alt)
        _c(ws,f"C{er}",it.name     if it else "",align="left",  border=True,wrap=True,bg=alt)
        _c(ws,f"D{er}",it.unit     if it else "",align="center",border=True,bg=alt)
        _c(ws,f"E{er}",it.quantity if it else "",align="center",border=True,bg=alt)
    sr=DR+25; ws.row_dimensions[sr].height=22
    _c(ws,f"A{sr}","Отпустил:",bold=True)
    ws.merge_cells(f"B{sr}:C{sr}"); _c(ws,f"B{sr}","______________________")
    _c(ws,f"D{sr}","Получил:",bold=True); _c(ws,f"E{sr}","______________")
    ws.print_area=f"A1:E{sr}"; ws.page_setup.orientation="portrait"
    ws.page_setup.fitToWidth=1; ws.sheet_properties.pageSetUpPr.fitToPage=True
    wb.save(path); return path

# ══════════════════════════════════════════════════════════════
# WORD HUJJAT
# ══════════════════════════════════════════════════════════════

def _tc(cell, text, bold=False, sz=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=""
    p=cell.paragraphs[0]; p.alignment=align
    r=p.add_run(str(text) if text else "")
    r.bold=bold; r.font.size=Pt(sz); r.font.name="Arial"

def build_word(nak, org_name, path, tpl=1):
    doc=Document(); s=doc.sections[0]
    s.left_margin=s.right_margin=s.top_margin=s.bottom_margin=Cm(1.5)
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("НАКЛАДНАЯ"); r.bold=True; r.font.size=Pt(18)
    s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    s2.add_run("на перемещение товаров, материалов").font.size=Pt(11)
    date_s=nak.created_at.strftime("%d.%m.%Y") if nak.created_at else ""
    m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    ri=m.add_run(f"№ {nak.number}   от {date_s} г."); ri.italic=True; ri.font.size=Pt(10)
    doc.add_paragraph(f"Организация: {org_name}").runs[0].bold=True
    doc.add_paragraph(f"Объект: {nak.object_text}")
    doc.add_paragraph(f"Отправитель: {nak.sender or '________________________'}")
    doc.add_paragraph(f"Получатель: {nak.receiver or '________________________'}")
    veh=f"{nak.vehicle_model} {nak.vehicle_plate}".strip() or "________________________"
    doc.add_paragraph(f"Машина: {veh}   Водитель: {nak.driver or '________________________'}")
    if nak.destination: doc.add_paragraph(f"Куда: {nak.destination}")
    n=max(len(nak.items),1)
    tbl=doc.add_table(rows=n+1,cols=5); tbl.style="Table Grid"
    tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["№","Бух.счет","Наименование","Ед. изм.","Кол-во"]):
        _tc(tbl.rows[0].cells[i],h,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx,it in enumerate(nak.items,1):
        row=tbl.rows[idx].cells
        _tc(row[0],idx,align=WD_ALIGN_PARAGRAPH.CENTER)
        _tc(row[1],it.code); _tc(row[2],it.name)
        _tc(row[3],it.unit,align=WD_ALIGN_PARAGRAPH.CENTER)
        _tc(row[4],it.quantity,align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    sg=doc.add_table(rows=1,cols=2)
    _tc(sg.rows[0].cells[0],"Отпустил: ____________________",bold=True)
    _tc(sg.rows[0].cells[1],"Получил: ____________________",bold=True)
    doc.save(path); return path

# ══════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════

async def to_pdf(docx_path, out_dir):
    prof=os.path.join(tempfile.gettempdir(),f"lo_{uuid.uuid4().hex}")
    os.makedirs(prof,exist_ok=True)
    proc=await asyncio.create_subprocess_exec(
        "soffice","--headless","--norestore",f"-env:UserInstallation=file://{prof}",
        "--convert-to","pdf","--outdir",out_dir,docx_path,
        stdout=asyncio.subprocess.PIPE,stderr=asyncio.subprocess.PIPE)
    try: _,err=await asyncio.wait_for(proc.communicate(),timeout=60)
    except asyncio.TimeoutError: proc.kill(); await proc.wait(); raise RuntimeError("PDF vaqt oshdi.")
    if proc.returncode!=0: raise RuntimeError(f"LibreOffice xato: {err.decode(errors='ignore')}")
    base=os.path.splitext(os.path.basename(docx_path))[0]
    pdf=os.path.join(out_dir,f"{base}.pdf")
    if not os.path.exists(pdf): raise RuntimeError("PDF yaratilmadi.")
    return pdf

# ══════════════════════════════════════════════════════════════
# PREVIEW PNG
# ══════════════════════════════════════════════════════════════

def build_preview(nak, org_name):
    W=900; M=30; RH=28; MAX=min(len(nak.items),24)
    H=220+len(org_name)//50*18+(MAX+2)*RH+80
    img=Image.new("RGB",(W,H),"white"); draw=ImageDraw.Draw(img)
    try:
        FP="/usr/share/fonts/truetype/dejavu/"
        fb=ImageFont.truetype(FP+"DejaVuSans-Bold.ttf",20)
        fh=ImageFont.truetype(FP+"DejaVuSans-Bold.ttf",12)
        fn=ImageFont.truetype(FP+"DejaVuSans.ttf",11)
    except: fb=fh=fn=ImageFont.load_default()
    y=M
    date_s=nak.created_at.strftime("%d.%m.%Y") if nak.created_at else ""
    draw.text((M,y),f"№ {nak.number}",font=fh,fill="#333")
    draw.text((W-180,y),date_s,font=fh,fill="#333"); y+=30
    title="НАКЛАДНАЯ"
    bbox=draw.textbbox((0,0),title,font=fb); tw=bbox[2]-bbox[0]
    draw.text(((W-tw)//2,y),title,font=fb,fill="#1F4E79"); y+=34
    sub="на перемещение товаров, материалов"
    bbox2=draw.textbbox((0,0),sub,font=fn); sw=bbox2[2]-bbox2[0]
    draw.text(((W-sw)//2,y),sub,font=fn,fill="#555"); y+=24
    draw.line([(M,y),(W-M,y)],fill="#1F4E79",width=2); y+=10
    for line in [
        f"Организация: {org_name[:70]}",
        f"Объект: {nak.object_text[:80]}",
        f"Отправитель: {nak.sender or '________________________'}",
        f"Получатель: {nak.receiver or '________________________'}",
    ]:
        draw.text((M,y),line,font=fn,fill="#222"); y+=20
    if nak.vehicle_plate or nak.driver:
        veh=f"{nak.vehicle_model} {nak.vehicle_plate}".strip()
        draw.text((M,y),f"Машина: {veh}   Водитель: {nak.driver or '___'}",font=fn,fill="#222"); y+=20
    y+=8
    CX=[M,M+45,M+110,M+490,M+600,W-M]
    draw.rectangle([(M,y),(W-M,y+RH)],fill="#D6E4F0")
    for i,h in enumerate(["№","Код","Наименование","Ед.","Кол-во"]):
        draw.text((CX[i]+4,y+7),h,font=fh,fill="#1F4E79")
    y+=RH
    for idx in range(MAX):
        bg="#FAFCFF" if idx%2==0 else "white"
        draw.rectangle([(M,y),(W-M,y+RH)],fill=bg)
        it=nak.items[idx]
        for i,v in enumerate([str(it.row_no),it.code,
                               it.name[:48] if len(it.name)>48 else it.name,
                               it.unit,it.quantity]):
            draw.text((CX[i]+4,y+7),v,font=fn,fill="#222")
        y+=RH
    draw.rectangle([(M-1,M-1),(W-M+1,H-M+1)],outline="#1F4E79",width=2)
    y+=10
    draw.text((M,y),"Отпустил: ____________________________",font=fn,fill="#222")
    draw.text((W//2,y),"Получил: ____________________________",font=fn,fill="#222")
    buf=io.BytesIO(); img.save(buf,format="PNG",optimize=True)
    return buf.getvalue()

# ══════════════════════════════════════════════════════════════
# GENERATE (barcha formatlar)
# ══════════════════════════════════════════════════════════════

async def generate_doc(nak, org_name, fmt, cfg, tpl=1):
    safe=nak.number.replace("/","-")
    if fmt=="excel": return build_excel(nak,org_name,os.path.join(GEN,f"nak_{safe}.xlsx"),tpl)
    docx=os.path.join(GEN,f"nak_{safe}.docx")
    build_word(nak,org_name,docx,tpl)
    if fmt=="word": return docx
    if fmt=="pdf":
        if cfg and not cfg.allow_pdf: return docx
        return await to_pdf(docx,GEN)
    return build_excel(nak,org_name,os.path.join(GEN,f"nak_{safe}.xlsx"),tpl)

# ══════════════════════════════════════════════════════════════
# GEMINI AI
# ══════════════════════════════════════════════════════════════

def _gemini():
    from database import settings
    import google.generativeai as genai
    if not settings.GEMINI_API_KEY: raise ValueError("GEMINI_API_KEY sozlanmagan")
    genai.configure(api_key=settings.GEMINI_API_KEY)
    return genai.GenerativeModel("gemini-2.5-flash")

def _clean_json(raw):
    raw=raw.strip()
    if raw.startswith("```"):
        parts=raw.split("```"); raw=parts[1] if len(parts)>1 else raw
        if raw.startswith("json"): raw=raw[4:]
    return raw.strip()

async def ai_ocr(image_bytes):
    model=_gemini()
    img=Image.open(io.BytesIO(image_bytes))
    prompt=('Bu qolyozma nakladnoy rasmi. Faqat JSON qaytaring:\n'
            '{"object_text":"","sender":"","receiver":"",'
            '"items":[{"code":"","name":"","unit":"кг/шт/м","quantity":""}]}\n'
            'Miqdorlarni vergul bilan yoz: 21,70')
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content([prompt,img]))
    return json.loads(_clean_json(resp.text))

async def ai_calc(expr):
    model=_gemini()
    prompt=(f"Qurilish/ombor hisob-kitobi mutaxassisi sifatida hal qiling:\n{expr}\n"
            "Qisqa, aniq javob bering. O'zbekcha. Natijani katta harflar bilan ko'rsating.")
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content(prompt))
    return resp.text.strip()

async def ai_search(query, materials):
    model=_gemini()
    mat_list="\n".join(f"{m['code']} | {m['name']} | {m['unit']}" for m in materials[:100])
    prompt=(f"'{query}' ni izlayapti. Quyidagi ro'yxatdan eng mos kodlarni top:\n{mat_list}\n"
            'Faqat JSON massiv: ["22551","9375"]')
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content(prompt))
    codes=json.loads(_clean_json(resp.text))
    return [m for m in materials if m.get("code") in codes][:8]

async def ai_report(stats):
    model=_gemini()
    prompt=(f"Omborxona hisoboti:\n{json.dumps(stats,ensure_ascii=False,indent=2)}\n"
            "O'zbekcha qisqa tahlil yozing (5-8 qator, emoji bilan).")
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content(prompt))
    return resp.text.strip()

async def ai_anomaly(user_name, actions):
    if not actions: return None
    model=_gemini()
    acts="\n".join(f"- {a.get('time','')} | {a.get('action','')} | {a.get('detail','')}" for a in actions[-20:])
    prompt=(f"Xodim: {user_name}\nHarakatlari:\n{acts}\n"
            "Shubhali holat bormi? (tun ishi, 10+/soat, takror chiqim)\n"
            "Borsa 1-2 gapda o'zbekcha yoz, bo'lmasa faqat OK yoz.")
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content(prompt))
    r=resp.text.strip()
    return None if r.upper()=="OK" else r

async def ai_briefing(stats):
    model=_gemini()
    prompt=(f"Kunlik brifing:\n{json.dumps(stats,ensure_ascii=False,indent=2)}\n"
            "Admin uchun 8-10 qatorli o'zbekcha xabar yozing (emoji bilan).")
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content(prompt))
    return resp.text.strip()

async def ai_suggest_template(items):
    model=_gemini()
    items_text="\n".join(f"- {it.get('name','')} ({it.get('unit','')})" for it in items[:10])
    tpls="1=Klassik,2=Ixcham,3=Korporativ,4=Batafsil,5=Sodda,6=Mahsulot,7=Transport,8=Ikki tilli,9=A5,10=Qurilish"
    prompt=f"Mahsulotlar:\n{items_text}\n\nShablonlar: {tpls}\n\nFaqat raqam qaytaring (1-10):"
    loop=asyncio.get_event_loop()
    resp=await loop.run_in_executor(None,lambda:model.generate_content(prompt))
    try: return int(resp.text.strip())
    except: return 1
